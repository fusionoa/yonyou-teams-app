from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

def set_run_font(run, font_name='微軟正黑體', size=11, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    try:
        run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font_name)
    except:
        pass

with open(r'C:\QClaw_Backup\TeamsApp\Yonyou_TeamsApp_安裝與架構手冊.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

in_code = False
code_block = []
table = None
prev_parts = []

for line in lines:
    line = line.rstrip()
    
    if line.startswith('---'):
        continue
    
    if line.startswith('# '):
        p = doc.add_heading(line[2:], level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=3)
    elif line.startswith('```'):
        in_code = not in_code
        if not in_code and code_block:
            p = doc.add_paragraph()
            p.style = 'No Spacing'
            for cb_line in code_block:
                run = p.add_run(cb_line + '\n')
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            code_block = []
    elif in_code:
        code_block.append(line)
    elif line.startswith('|') and '---' not in line:
        parts = [c.strip() for c in line.split('|')[1:-1]]
        if len(parts) > 0 and parts[0]:
            table = doc.add_table(rows=1, cols=len(parts))
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, cell in enumerate(parts):
                hdr[i].text = cell
                for p in hdr[i].paragraphs:
                    for run in p.runs:
                        set_run_font(run)
            prev_parts = parts
        elif len(parts) == 0 or parts[0] == '':
            if prev_parts:
                row = table.add_row().cells
                for i in range(len(prev_parts)):
                    row[i].text = ''
    elif line.startswith('- [ ] ') or line.startswith('- [x] '):
        doc.add_paragraph(line[6:], style='List Bullet')
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('> '):
        p = doc.add_paragraph(line[2:])
        p.paragraph_format.left_indent = Inches(0.3)
    elif line.strip() == '':
        doc.add_paragraph()
    else:
        if '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.strip('*'))
                    run.font.bold = True
                else:
                    p.add_run(part)
        else:
            doc.add_paragraph(line)

doc.save(r'C:\QClaw_Backup\TeamsApp\Yonyou_TeamsApp_安裝與架構手冊.docx')
print('Done')
